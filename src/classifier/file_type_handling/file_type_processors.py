# Processors for different file types

import PyPDF2
from PIL import Image
import openpyxl
from docx import Document
from pathlib import Path
from typing import Dict, Any
from werkzeug.datastructures import FileStorage
from src.classifier.categories.category_loader import get_categories_for_industry
from src.llm_call import classify_with_llm
from src.classifier.file_type_handling.document_utils import (
    pdf_to_image,
    excel_to_image,
    word_to_image,
    optimize_image_for_llm,
    ocr_extract_text,
    extract_excel_text
)

def classify_file(file: FileStorage, industry: str = 'finance') -> dict:
    """Main classifier - simple routing to processors"""
    
    if not file.filename:
        return {'classification': 'no_filename', 'confidence': 0.0}
    
    file_extension = Path(file.filename).suffix.lower().lstrip('.')
    
    # Route to appropriate processor
    processors = {
        'pdf': classify_pdf,
        'jpg': classify_image,
        'jpeg': classify_image,
        'png': classify_image,
        'gif': classify_image,
        'bmp': classify_image,
        'tiff': classify_image,
        'xls': classify_excel,
        'xlsx': classify_excel,
        'doc': classify_word,
        'docx': classify_word,
    }
    
    processor = processors.get(file_extension)
    if not processor:
        return {
            'classification': 'unsupported_file_type',
            'confidence': 0.0,
            'metadata': {'filename': file.filename, 'file_type': file_extension},
            'industry': industry
        }
    
    return processor(file, industry)

def classify_excel(file: FileStorage, industry: str = 'finance') -> dict:
    """
    Excel processor
    """
    file.seek(0)
    
    try:
        workbook = openpyxl.load_workbook(file, read_only=True)

        # Get the file size
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        sheet_count = len(workbook.worksheets)
        sheet_names = [sheet.title for sheet in workbook.worksheets]
        first_sheet = workbook.active

        # Metadata dictionary
        metadata = {
            'filename': file.filename,
            'file_type': 'excel',
            'file_size': file_size,
            'page_count': sheet_count,
            'sheet_names': sheet_names[:3],  # First 3 sheet names
            'max_rows': first_sheet.max_row,
            'max_columns': first_sheet.max_column,
        }

        # Process content
        image_data = excel_to_image(file, sheet=0)
        text_content = extract_excel_text(file, sheet=0)

        # Call LLM function
        result = classify_with_llm(
            image_data=image_data,
            text_content=text_content,
            metadata=metadata,
            industry=industry
        )

        # Add processor-specific info
        result['metadata'] = metadata
        result['industry'] = industry
        
        return result
    
    except Exception as e:
        return {
            'classification': 'processing_error',
            'confidence': 0.0,
            'reasoning': f"Excel processing failed: {str(e)}",
            'metadata': {'filename': file.filename, 'file_type': 'excel'},
            'industry': industry
        }

def classify_word(file: FileStorage, industry: str = 'finance') -> dict:
    """Word processor"""
    
    file.seek(0)
    
    try:
        doc = Document(file)

        # Get the file size
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        paragraph_count = len(doc.paragraphs)
        text_content = '\n'.join([p.text for p in doc.paragraphs[:3]])  # First 3 paragraphs
        estimated_pages = max(1, len(text_content) // 2000) # Rough estimate of pages

        # Metadata dictionary
        metadata = {
            'filename': file.filename,
            'file_type': 'word',
            'file_size': file_size,
            'page_count': estimated_pages,
            'paragraph_count': paragraph_count,
            'character_count': len(text_content),

        }

        # Process content
        image_data = word_to_image(file, page=0)

        # Call LLM function
        result = classify_with_llm(
            image_data=image_data,
            text_content=text_content,
            metadata=metadata,
            industry=industry
        )

        # Add processor-specific info
        result['metadata'] = metadata
        result['industry'] = industry

        return result

    except Exception as e:
        return {
            'classification': 'processing_error',
            'confidence': 0.0,
            'reasoning': f"Word processing failed: {str(e)}",
            'metadata': {'filename': file.filename, 'file_type': 'word'},
            'industry': industry
        }
    
def classify_image(file: FileStorage, industry: str = 'finance') -> dict:
    """
    Image processor
    """
    
    file.seek(0)
    
    try:
        with Image.open(file) as img:
            file.seek(0, 2)
            file_size = file.tell()
            file.seek(0)
            
            # Metadata dictionary
            metadata = {
                'filename': file.filename,
                'file_type': 'image',
                'file_size': file_size,
                'width': img.width,
                'height': img.height,
                'mode': img.mode,
                'page_count': 1,  # Single image
                'format': img.format
            }
            
            # Process content
            optimized_image = optimize_image_for_llm(file)
            text_content = ocr_extract_text(file)
            
            # Call LLM function
            result = classify_with_llm(
                image_data=optimized_image,
                text_content=text_content,
                metadata=metadata,
                industry=industry
            )
            
            # Add processor-specific info
            result['metadata'] = metadata
            result['industry'] = industry
            
            return result
            
    except Exception as e:
        return {
            'classification': 'processing_error',
            'confidence': 0.0,
            'reasoning': f"Image processing failed: {str(e)}",
            'metadata': {'filename': file.filename, 'file_type': 'image'},
            'industry': industry
        }

def classify_pdf(file: FileStorage, industry: str = 'finance') -> dict:
    """
    PDF processor
    """
    
    file.seek(0)
    
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        page_count = len(pdf_reader.pages)
        
        # Extract text
        first_page_text = ""
        has_text = False
        if page_count > 0:
            first_page_text = pdf_reader.pages[0].extract_text()
            has_text = len(first_page_text.strip()) > 10
        
        # Get file size
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)
        
        # Build metadata dictionary
        metadata = {
            'filename': file.filename,
            'file_type': 'pdf',
            'file_size': file_size,
            'page_count': page_count,
            'has_text': has_text,
            'first_page_text': first_page_text[:500]  # First 500 characters
        }
        
        # Process content
        image_data = pdf_to_image(file, page=0)
        
        # Call LLM
        result = classify_with_llm(
            image_data=image_data,
            text_content=first_page_text,
            metadata=metadata,
            industry=industry
        )
        
        # Add processor-specific info
        result['metadata'] = metadata
        result['industry'] = industry
        
        return result
        
    except Exception as e:
        return {
            'classification': 'processing_error',
            'confidence': 0.0,
            'reasoning': f"PDF processing failed: {str(e)}",
            'metadata': {'filename': file.filename, 'file_type': 'pdf'},
            'industry': industry
        }
